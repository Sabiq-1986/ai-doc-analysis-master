"""
Word Parser - Full .docx extraction with XML-level features.
Heading styles, headers/footers per section, footnotes/endnotes,
comments/tracked changes, embedded images with OCR, hyperlinks, lists, bold detection.
"""
import io
import logging
import os
import re
import zipfile
from pathlib import Path
from typing import List, Optional
from xml.etree import ElementTree

from langchain_core.documents import Document

logger = logging.getLogger(__name__)

# python-docx
try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# olefile for legacy .doc
try:
    import olefile
    OLEFILE_AVAILABLE = True
except ImportError:
    OLEFILE_AVAILABLE = False

# Namespace map for Word XML
WORD_NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
}


class WordParser:
    """Full python-docx extraction with XML-level features."""

    def __init__(self, image_parser=None):
        self.image_parser = image_parser

    def parse(self, file_path: str) -> List[Document]:
        """Parse Word document (.docx or .doc)."""
        ext = Path(file_path).suffix.lower()
        filename = Path(file_path).name

        if ext == '.docx':
            return self._parse_docx(file_path, filename)
        elif ext == '.doc':
            return self._parse_doc(file_path, filename)
        else:
            logger.warning(f"[WORD] Unsupported extension: {ext}")
            return []

    def _parse_docx(self, file_path: str, filename: str) -> List[Document]:
        """Parse .docx file with comprehensive extraction."""
        if not DOCX_AVAILABLE:
            logger.error("[WORD] python-docx not installed")
            return []

        documents = []

        try:
            doc = DocxDocument(file_path)

            # Document properties
            props = self._extract_properties(doc)
            if props:
                documents.append(Document(
                    page_content=props,
                    metadata={"source": filename, "type": "docx", "section": "properties"}
                ))

            # Main body text with headings and bold detection
            body_docs = self._extract_body(doc, filename)
            documents.extend(body_docs)

            # Tables
            table_docs = self._extract_tables(doc, filename)
            documents.extend(table_docs)

            # Headers and footers
            hf_docs = self._extract_headers_footers(doc, filename)
            documents.extend(hf_docs)

            # Footnotes and endnotes (via XML)
            fn_docs = self._extract_footnotes_endnotes(file_path, filename)
            documents.extend(fn_docs)

            # Comments (via XML)
            comment_docs = self._extract_comments(file_path, filename)
            documents.extend(comment_docs)

            # Hyperlinks
            link_docs = self._extract_hyperlinks(doc, filename)
            documents.extend(link_docs)

            # Embedded images
            image_docs = self._extract_images(doc, file_path, filename)
            documents.extend(image_docs)

        except Exception as e:
            logger.error(f"[WORD] Error parsing {filename}: {e}")
            import traceback
            traceback.print_exc()

        return documents

    def _extract_properties(self, doc) -> Optional[str]:
        """Extract document properties."""
        try:
            props = doc.core_properties
            lines = ["[Document Properties]"]

            if props.title:
                lines.append(f"Title: {props.title}")
            if props.author:
                lines.append(f"Author: {props.author}")
            if props.subject:
                lines.append(f"Subject: {props.subject}")
            if props.keywords:
                lines.append(f"Keywords: {props.keywords}")
            if props.created:
                lines.append(f"Created: {props.created}")
            if props.modified:
                lines.append(f"Modified: {props.modified}")
            if props.last_modified_by:
                lines.append(f"Last Modified By: {props.last_modified_by}")
            if props.revision:
                lines.append(f"Revision: {props.revision}")
            if props.category:
                lines.append(f"Category: {props.category}")
            if props.comments:
                lines.append(f"Description: {props.comments}")

            if len(lines) > 1:
                return '\n'.join(lines)
        except Exception:
            pass
        return None

    def _extract_body(self, doc, filename: str) -> List[Document]:
        """Extract body text with heading detection and bold emphasis."""
        documents = []
        current_section = []
        current_heading = None
        para_idx = 0

        for para in doc.paragraphs:
            para_idx += 1
            text = para.text.strip()
            if not text:
                continue

            # Detect heading
            style_name = para.style.name.lower() if para.style else ""
            is_heading = 'heading' in style_name or style_name.startswith('title')

            # Detect bold emphasis
            has_bold = any(run.bold for run in para.runs if run.text.strip())

            if is_heading:
                # Save previous section
                if current_section:
                    section_text = '\n'.join(current_section)
                    metadata = {
                        "source": filename, "type": "docx",
                        "section": current_heading or "body"
                    }
                    documents.append(Document(page_content=section_text, metadata=metadata))
                    current_section = []

                # Extract heading level
                heading_level = ""
                for char in style_name:
                    if char.isdigit():
                        heading_level += char

                prefix = f"{'#' * int(heading_level)} " if heading_level else "# "
                current_heading = text
                current_section.append(f"{prefix}{text}")
            else:
                # Apply bold emphasis markers
                if has_bold and len(para.runs) > 1:
                    parts = []
                    for run in para.runs:
                        if run.text.strip():
                            if run.bold:
                                parts.append(f"**{run.text}**")
                            else:
                                parts.append(run.text)
                    text = ''.join(parts)

                current_section.append(text)

        # Don't forget last section
        if current_section:
            section_text = '\n'.join(current_section)
            documents.append(Document(
                page_content=section_text,
                metadata={"source": filename, "type": "docx",
                          "section": current_heading or "body"}
            ))

        return documents

    def _extract_tables(self, doc, filename: str) -> List[Document]:
        """Extract tables with cell content."""
        documents = []

        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    text = cell.text.strip()
                    cells.append(text)
                if any(cells):
                    rows.append(' | '.join(cells))

            if rows:
                table_text = f"[Table {t_idx + 1}]\n" + '\n'.join(rows)
                documents.append(Document(
                    page_content=table_text,
                    metadata={"source": filename, "type": "docx",
                              "section": "table", "table_index": t_idx + 1}
                ))

        return documents

    def _extract_headers_footers(self, doc, filename: str) -> List[Document]:
        """Extract headers and footers per section."""
        documents = []

        for s_idx, section in enumerate(doc.sections):
            # Header
            try:
                header = section.header
                if header and header.paragraphs:
                    text = '\n'.join(p.text.strip() for p in header.paragraphs if p.text.strip())
                    if text:
                        documents.append(Document(
                            page_content=f"[Header - Section {s_idx+1}]\n{text}",
                            metadata={"source": filename, "type": "docx",
                                      "section": "header", "section_index": s_idx + 1}
                        ))
            except Exception:
                pass

            # Footer
            try:
                footer = section.footer
                if footer and footer.paragraphs:
                    text = '\n'.join(p.text.strip() for p in footer.paragraphs if p.text.strip())
                    if text:
                        documents.append(Document(
                            page_content=f"[Footer - Section {s_idx+1}]\n{text}",
                            metadata={"source": filename, "type": "docx",
                                      "section": "footer", "section_index": s_idx + 1}
                        ))
            except Exception:
                pass

        return documents

    def _extract_footnotes_endnotes(self, file_path: str, filename: str) -> List[Document]:
        """Extract footnotes and endnotes from Word XML."""
        documents = []

        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                # Footnotes
                if 'word/footnotes.xml' in zf.namelist():
                    xml_data = zf.read('word/footnotes.xml')
                    root = ElementTree.fromstring(xml_data)
                    notes = []

                    for fn in root.findall('.//w:footnote', WORD_NS):
                        fn_id = fn.get(f'{{{WORD_NS["w"]}}}id', '')
                        if fn_id in ('-1', '0'):  # Skip separator footnotes
                            continue
                        text_parts = []
                        for para in fn.findall('.//w:p', WORD_NS):
                            for run in para.findall('.//w:r', WORD_NS):
                                t = run.find('w:t', WORD_NS)
                                if t is not None and t.text:
                                    text_parts.append(t.text)
                        text = ''.join(text_parts).strip()
                        if text:
                            notes.append(f"[{fn_id}] {text}")

                    if notes:
                        documents.append(Document(
                            page_content="[Footnotes]\n" + '\n'.join(notes),
                            metadata={"source": filename, "type": "docx", "section": "footnotes"}
                        ))

                # Endnotes
                if 'word/endnotes.xml' in zf.namelist():
                    xml_data = zf.read('word/endnotes.xml')
                    root = ElementTree.fromstring(xml_data)
                    notes = []

                    for en in root.findall('.//w:endnote', WORD_NS):
                        en_id = en.get(f'{{{WORD_NS["w"]}}}id', '')
                        if en_id in ('-1', '0'):
                            continue
                        text_parts = []
                        for para in en.findall('.//w:p', WORD_NS):
                            for run in para.findall('.//w:r', WORD_NS):
                                t = run.find('w:t', WORD_NS)
                                if t is not None and t.text:
                                    text_parts.append(t.text)
                        text = ''.join(text_parts).strip()
                        if text:
                            notes.append(f"[{en_id}] {text}")

                    if notes:
                        documents.append(Document(
                            page_content="[Endnotes]\n" + '\n'.join(notes),
                            metadata={"source": filename, "type": "docx", "section": "endnotes"}
                        ))

        except Exception as e:
            logger.debug(f"[WORD] Footnote/endnote extraction failed: {e}")

        return documents

    def _extract_comments(self, file_path: str, filename: str) -> List[Document]:
        """Extract comments from Word XML."""
        documents = []

        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                if 'word/comments.xml' not in zf.namelist():
                    return []

                xml_data = zf.read('word/comments.xml')
                root = ElementTree.fromstring(xml_data)
                comments = []

                for comment in root.findall('.//w:comment', WORD_NS):
                    author = comment.get(f'{{{WORD_NS["w"]}}}author', 'Unknown')
                    date = comment.get(f'{{{WORD_NS["w"]}}}date', '')
                    text_parts = []
                    for para in comment.findall('.//w:p', WORD_NS):
                        for run in para.findall('.//w:r', WORD_NS):
                            t = run.find('w:t', WORD_NS)
                            if t is not None and t.text:
                                text_parts.append(t.text)
                    text = ''.join(text_parts).strip()
                    if text:
                        date_str = f" ({date[:10]})" if date else ""
                        comments.append(f"{author}{date_str}: {text}")

                if comments:
                    documents.append(Document(
                        page_content="[Comments]\n" + '\n'.join(comments),
                        metadata={"source": filename, "type": "docx",
                                  "section": "comments", "comment_count": len(comments)}
                    ))

        except Exception as e:
            logger.debug(f"[WORD] Comment extraction failed: {e}")

        return documents

    def _extract_hyperlinks(self, doc, filename: str) -> List[Document]:
        """Extract hyperlinks from document."""
        links = []

        try:
            rels = doc.part.rels
            for para in doc.paragraphs:
                for run in para.runs:
                    # Check for hyperlink in XML
                    parent = run._element.getparent()
                    if parent is not None:
                        hyperlink = parent.find('.//w:hyperlink', WORD_NS)
                        if hyperlink is not None:
                            r_id = hyperlink.get(f'{{{WORD_NS["r"]}}}id', '')
                            if r_id and r_id in rels:
                                url = rels[r_id].target_ref
                                display = run.text.strip()
                                if url and display:
                                    links.append(f"{display} -> {url}")
        except Exception:
            pass

        # Also check inline hyperlinks in paragraphs
        try:
            for para in doc.paragraphs:
                elem = para._element
                for hyperlink in elem.findall('.//w:hyperlink', WORD_NS):
                    r_id = hyperlink.get(f'{{{WORD_NS["r"]}}}id', '')
                    text_parts = []
                    for run in hyperlink.findall('.//w:r', WORD_NS):
                        t = run.find('w:t', WORD_NS)
                        if t is not None and t.text:
                            text_parts.append(t.text)
                    display = ''.join(text_parts).strip()

                    if r_id and r_id in doc.part.rels:
                        url = doc.part.rels[r_id].target_ref
                        if url and display and f"{display} -> {url}" not in links:
                            links.append(f"{display} -> {url}")
        except Exception:
            pass

        if links:
            # Deduplicate
            unique_links = list(dict.fromkeys(links))
            return [Document(
                page_content="[Hyperlinks]\n" + '\n'.join(unique_links),
                metadata={"source": filename, "type": "docx",
                          "section": "hyperlinks", "link_count": len(unique_links)}
            )]
        return []

    def _extract_images(self, doc, file_path: str, filename: str) -> List[Document]:
        """Extract and OCR embedded images."""
        if not self.image_parser:
            return []

        documents = []

        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                image_files = [
                    f for f in zf.namelist()
                    if f.startswith('word/media/')
                    and f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.emf', '.wmf'))
                ]

                for img_idx, img_file in enumerate(image_files):
                    try:
                        img_data = zf.read(img_file)

                        if len(img_data) < 1000:
                            continue

                        # Skip EMF/WMF (vector formats)
                        if img_file.lower().endswith(('.emf', '.wmf')):
                            continue

                        result = self.image_parser.extract_text(
                            img_data,
                            filename=f"{filename}_img{img_idx+1}",
                            mode="auto"
                        )

                        text = result.get("text", "")
                        if text and len(text.strip()) > 10:
                            documents.append(Document(
                                page_content=text,
                                metadata={
                                    "source": filename,
                                    "type": "docx_image",
                                    "image_index": img_idx + 1,
                                    "image_file": Path(img_file).name,
                                }
                            ))

                    except Exception as e:
                        logger.debug(f"[WORD] Image {img_file} extraction failed: {e}")

        except Exception as e:
            logger.debug(f"[WORD] Image extraction failed: {e}")

        return documents

    def _parse_doc(self, file_path: str, filename: str) -> List[Document]:
        """Parse legacy .doc format using olefile."""
        if not OLEFILE_AVAILABLE:
            logger.warning("[WORD] olefile not installed, cannot parse .doc files")
            return []

        try:
            ole = olefile.OleFileIO(file_path)

            text = ""
            if ole.exists('WordDocument'):
                stream = ole.openstream('WordDocument')
                data = stream.read()
                # Extract text from binary stream
                text_parts = re.findall(b'[\x20-\x7e\xc0-\xff]+', data)
                text = ' '.join(part.decode('latin-1', errors='replace') for part in text_parts)

            ole.close()

            if text.strip():
                return [Document(
                    page_content=text.strip(),
                    metadata={"source": filename, "type": "doc",
                              "extraction_mode": "olefile"}
                )]

        except Exception as e:
            logger.error(f"[WORD] .doc parsing failed: {e}")

        return []
